#docx file manipulation
# page and paragraph manipulation:

import docx

doc = docx.Document()

doc.add_heading('VVCE - Autonomous Institution', 0)

doc.add_heading('Normal Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='Normal')


doc.add_heading('Caption Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='Caption')

doc.add_heading('Title Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='Title')

doc.add_heading('Heading 1 Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='Heading 1')

doc.add_heading('Macro Text Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='macro')

doc.add_page_break() 

doc.add_heading('Quoted Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='Quote')

doc.add_heading('TOC Heading Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='TOC Heading')

doc.add_heading('Subtitle Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='Subtitle')

doc.add_heading('No Spacing Style:', 3)
doc.add_paragraph('VVCE is an autonomous institution.', style='No Spacing')
doc.add_paragraph('VVCE is affiliated to vtu.')
doc.save('abapython.docx')


#paragraph identificatin and printing :

from docx import Document

def get_paragraphs():
file_path = input("Enter the file path of the Word file: ")
    doc = Document(file_path)

    for para in doc.paragraphs:
        print(para.text)
print()

get_paragraphs()



# conversion to upper case letters:

from docx import Document


document = Document('file-sample_1MB.docx')


for paragraph in document.paragraphs:

paragraph.text = paragraph.text.upper()


for table in document.tables:
    for row in table.rows:
        for cell in row.cells:

cell.text = cell.text.upper()


document.save('converted_document.docx')
